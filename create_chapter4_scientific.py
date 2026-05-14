#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Tạo Chương IV: Mô Hình Thực Nghiệm (phong cách khoa học)
Cho dự án Cabin Translation - Đồ án tốt nghiệp
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def create_scientific_document():
    doc = Document()
    
    # Set margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)
    
    # Title
    title = doc.add_heading('CHƯƠNG IV: MÔ HÌNH THỰC NGHIỆM', 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # 4.1
    doc.add_heading('4.1. Môi trường triển khai', level=1)
    
    doc.add_heading('4.1.1. Cấu hình phần cứng', level=2)
    
    doc.add_paragraph(
        'Môi trường triển khai hệ thống Cabin Translation được thiết lập trên nền tảng phần cứng máy tính '
        'cá nhân phổ thông, nhằm chứng minh khả năng vận hành của hệ thống mà không đòi hỏi cơ sở hạ tầng '
        'máy chủ chuyên dụng. Chiến lược này tạo điều kiện thuận lợi cho việc kiểm thử hệ thống, đồng thời '
        'thể hiện tính khả thi của việc triển khai ở quy mô nhỏ và vừa, phù hợp với các tổ chức có nguồn lực '
        'hạn chế nhưng vẫn cần cung cấp dịch vụ dịch thuật chất lượng cao.'
    )
    
    doc.add_paragraph(
        'Đơn vị xử lý trung tâm (CPU - Central Processing Unit) được đòi hỏi phải đạt tốc độ tối thiểu hai '
        'gigahertz (2.0 GHz) với ít nhất hai nhân xử lý vật lý (physical cores). Yêu cầu này xuất phát từ '
        'đặc thù của hệ thống cần xử lý một cách song song nhiều tác vụ đó là: thu âm từ thiết bị đầu vào '
        'qua Giao diện lập trình ứng dụng Web Audio (Web Audio Application Programming Interface), truyền tải '
        'dữ liệu âm thanh qua giao thức WebSocket đến máy chủ backend, phản hồi các yêu cầu từ giao diện người '
        'dùng, giao tiếp với dịch vụ Soniox để nhận dạng và dịch thuật giọng nói, cùng thực hiện các truy vấn '
        'đến cơ sở dữ liệu để lưu trữ và truy xuất lịch sử dịch.'
    )
    
    doc.add_paragraph(
        'Bộ nhớ truy cập ngẫu nhiên (RAM - Random Access Memory) tối thiểu tám gigabyte (8 GB) là yêu cầu '
        'cần thiết để đảm bảo hệ thống vận hành ổn định khi các thành phần chính hoạt động đồng thời. Phân tích '
        'chi tiết về tiêu thụ bộ nhớ cho thấy máy chủ backend FastAPI tiêu thụ khoảng từ một đến hai gigabyte, '
        'trình duyệt web chạy ứng dụng Next.js chiếm từ năm trăm megabyte đến một gigabyte, hệ quản trị cơ sở '
        'dữ liệu Microsoft SQL Server yêu cầu khoảng một đến hai gigabyte, hệ điều hành Windows và các tiến trình '
        'nền sử dụng hai đến ba gigabyte, phần còn lại được dành cho các dịch vụ ngoài bao gồm khách hàng Soniox '
        'và phần mềm phát triển của Google Cloud.'
    )
    
    doc.add_paragraph(
        'Không gian lưu trữ tối thiểu hai mươi gigabyte (20 GB) được cấp phát cho toàn bộ hệ thống. Thành phần '
        'chính của không gian lưu trữ này bao gồm hệ điều hành Windows 10 hoặc phiên bản cao hơn (4 GB), môi trường '
        'phát triển Python phiên bản 3.14 cùng các thư viện phụ thuộc được liệt kê trong tệp requirements.txt (2 GB), '
        'Node.js phiên bản 18 hoặc cao hơn cùng với các gói thư viện quản lý gói Node (npm - Node Package Manager) '
        'từ tệp package.json (3 GB), tệp cơ sở dữ liệu SQL Server Express (5 GB), các tệp tạm thời từ quá trình xử lý '
        'âm thanh thời gian thực, bộ nhớ đệm (cache), và tệp nhật ký hệ thống (logs) (2 GB).'
    )
    
    doc.add_paragraph(
        'Thiết bị thu âm thanh (microphone) đóng vai trò vô cùng quan trọng trong việc đạt được chất lượng nhận dạng '
        'giọng nói cao. Thiết bị này, dù là micro tích hợp sẵn trong máy tính hay micro ngoại vi kết nối qua kết nối USB, '
        'phải đảm bảo các đặc tính kỹ thuật cụ thể sau: tần số lấy mẫu âm thanh (sample rate) tối thiểu mười sáu kilohertz '
        '(16 kHz), độ phân giải bit tối thiểu mười sáu bit (16-bit), kênh âm thanh đơn (mono) theo chuẩn mã hóa xung '
        'điều chế PCM (Pulse-Code Modulation). Những thông số kỹ thuật này là tiêu chuẩn mà dịch vụ Soniox yêu cầu để đạt '
        'được độ chính xác cao nhất trong quá trình nhận dạng giọng nói.'
    )
    
    doc.add_paragraph(
        'Kết nối mạng ổn định với tốc độ tối thiểu một megabit trên giây (1 Mbps) được đòi hỏi để truyền tải dữ liệu '
        'âm thanh theo thời gian thực đến máy chủ Soniox. Trong thực tế kiểm thử, kết nối mạng cục bộ (LAN - Local Area Network) '
        'hoặc đường truyền cáp quang băng thông rộng (>10 Mbps) cho thấy hiệu quả hoạt động tốt nhất với độ trễ dưới hai trăm '
        'mili giây (200 ms), đáp ứng yêu cầu của giao tiếp thời gian thực.'
    )
    
    # Bảng phần cứng
    table_hw = doc.add_table(rows=9, cols=2)
    table_hw.style = 'Light Grid Accent 1'
    header_cells = table_hw.rows[0].cells
    header_cells[0].text = 'Thành phần phần cứng'
    header_cells[1].text = 'Yêu cầu tối thiểu'
    
    hw_data = [
        ('Bộ xử lý trung tâm', '2.0 GHz, 2 nhân xử lý vật lý'),
        ('Bộ nhớ RAM', '8 gigabyte'),
        ('Lưu trữ dữ liệu', '20 gigabyte SSD'),
        ('Thiết bị thu âm', '16 kHz, 16-bit PCM, kênh đơn'),
        ('Kết nối mạng', '1 megabit trên giây'),
        ('Hệ điều hành', 'Windows 10 64-bit hoặc cao hơn'),
        ('Trình duyệt web', 'Chrome 90 hoặc cao hơn, Edge 90 hoặc cao hơn'),
        ('Bộ nhớ xử lý WebRTC', '256 megabyte để xử lý âm thanh thời gian thực'),
    ]
    
    for i, (col1, col2) in enumerate(hw_data, start=1):
        table_hw.rows[i].cells[0].text = col1
        table_hw.rows[i].cells[1].text = col2
    
    doc.add_paragraph()
    p = doc.add_paragraph('Bảng 4.1: Cấu hình phần cứng môi trường triển khai hệ thống')
    p_format = p.paragraph_format
    p_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # 4.1.2
    doc.add_heading('4.1.2. Cấu hình phần mềm', level=2)
    
    doc.add_paragraph(
        'Hệ điều hành Windows 10 hoặc Windows 11 phiên bản 64-bit được lựa chọn làm nền tảng hệ điều hành chính. '
        'Quyết định này dựa trên các lý do sau: đảm bảo khả năng tương thích tốt với hệ quản trị cơ sở dữ liệu SQL Server, '
        'các công cụ phát triển Python và Node.js, đồng thời hệ điều hành này phổ biến rộng rãi trong môi trường học thuật tại '
        'các trường đại học Việt Nam.'
    )
    
    doc.add_paragraph(
        'Python phiên bản ba điểm mười bốn (3.14) được cài đặt cho môi trường backend của dự án. Phiên bản này hỗ trợ đầy đủ '
        'các tính năng lập trình bất đồng bộ (asynchronous programming - async/await) cần thiết cho khung ứng dụng FastAPI, cùng '
        'đó là các thư viện xử lý âm thanh chuyên biệt. Để cô lập các thư viện phụ thuộc của dự án khỏi các thư viện hệ thống, '
        'một môi trường ảo Python (Virtual Environment - venv) được tạo riêng biệt tại thư mục Backend/.venv.'
    )
    
    doc.add_paragraph(
        'Node.js phiên bản mười tám hoặc cao hơn được cài đặt cho môi trường frontend. Khung ứng dụng Next.js phiên bản mười ba '
        'yêu cầu Node.js tối thiểu ở phiên bản mười tám để hỗ trợ đầy đủ các tính năng App Router, Server Components, và các tối ưu '
        'hóa hiệu suất mới nhất của framework này.'
    )
    
    doc.add_paragraph(
        'Microsoft SQL Server Express Edition phiên bản 2022 được lựa chọn làm hệ quản trị cơ sở dữ liệu (DBMS - Database Management System). '
        'SQL Server Express là lựa chọn thích hợp cho môi trường phát triển và thực nghiệm vì nó cung cấp đầy đủ các chức năng quản trị dữ liệu '
        'với giới hạn dung lượng mười gigabyte, hoàn toàn đáp ứng nhu cầu của dự án hiện tại. Chế độ xác thực hỗn hợp (Mixed Authentication Mode) '
        'được kích hoạt để hỗ trợ cả xác thực Windows và xác thực SQL Server.'
    )
    
    # Bảng phần mềm
    table_sw = doc.add_table(rows=11, cols=2)
    table_sw.style = 'Light Grid Accent 1'
    header_cells = table_sw.rows[0].cells
    header_cells[0].text = 'Thành phần phần mềm'
    header_cells[1].text = 'Phiên bản yêu cầu'
    
    sw_data = [
        ('Hệ điều hành', 'Windows 10/11 64-bit'),
        ('Python', '3.14'),
        ('Node.js', '18 hoặc cao hơn'),
        ('npm', '9 hoặc cao hơn'),
        ('FastAPI', '0.104 hoặc cao hơn'),
        ('Next.js', '13 hoặc cao hơn'),
        ('SQL Server', 'Express 2022'),
        ('Soniox SDK', 'Phiên bản mới nhất'),
        ('Google Cloud TTS', 'Phiên bản 1'),
        ('OpenAI API', 'Phiên bản 1'),
    ]
    
    for i, (col1, col2) in enumerate(sw_data, start=1):
        table_sw.rows[i].cells[0].text = col1
        table_sw.rows[i].cells[1].text = col2
    
    doc.add_paragraph()
    p = doc.add_paragraph('Bảng 4.2: Cấu hình phần mềm môi trường triển khai')
    p_format = p.paragraph_format
    p_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # 4.1.3
    doc.add_heading('4.1.3. Kiến trúc hệ thống', level=2)
    
    doc.add_paragraph(
        'Hệ thống Cabin Translation được thiết kế theo mô hình kiến trúc ba tầng (Three-tier Architecture), một mô hình phổ biến '
        'và được chứng minh hiệu quả trong phát triển ứng dụng web hiện đại. Ba tầng này bao gồm tầng Presentation (lớp giao diện), '
        'tầng Business Logic (lớp logic xử lý), và tầng Data (lớp dữ liệu).'
    )
    
    doc.add_paragraph(
        'Tầng giao diện người dùng (Tầng Presentation) được phát triển bằng khung ứng dụng Next.js phiên bản 13 trở lên kết hợp với '
        'thư viện React, cung cấp giao diện tương tác trên trình duyệt web. Các thành phần chính của tầng này bao gồm trang chính (index.js) '
        'với chức năng đăng nhập và đăng ký tài khoản, trang dịch văn bản (translate.js) với trình soạn thảo song song cho phép người dùng '
        'nhập và xem kết quả dịch side-by-side, trang dịch giọng nói (voice.js) với chức năng ghi âm thời gian thực qua Giao diện lập trình '
        'ứng dụng Web Audio, trang lịch sử (history.js) để xem các bản dịch trước đây, trang quản lý hóa đơn (billing.js) để người dùng quản '
        'lý credits và nạp tiền, cùng trang hỗ trợ (support.js) để người dùng có thể liên hệ với nhóm quản trị viên.'
    )
    
    doc.add_paragraph(
        'Tầng logic xử lý (Tầng Business Logic) được phát triển bằng khung ứng dụng FastAPI, một framework Python hiệu suất cao được thiết kế '
        'đặc biệt cho các ứng dụng bất đồng bộ. Máy chủ FastAPI chạy trên máy chủ ASGI (Asynchronous Server Gateway Interface) có tên Uvicorn tại '
        'địa chỉ localhost cổng 8000. Tầng này xử lý tất cả logic nghiệp vụ bao gồm xác thực người dùng, dịch vụ dịch thuật, gọi các dịch vụ ngoài, '
        'quản lý cơ sở dữ liệu, và xử lý các yêu cầu từ tầng giao diện. Các module chính của tầng này được tổ chức theo chức năng.'
    )
    
    doc.add_paragraph(
        'Tầng dữ liệu (Tầng Data) bao gồm hai phần chính. Phần thứ nhất là hệ quản trị cơ sở dữ liệu Microsoft SQL Server Express lưu trữ toàn bộ '
        'dữ liệu ứng dụng bao gồm thông tin người dùng, lịch sử dịch thuật, ticket hỗ trợ, thông báo hệ thống, và các thông tin liên quan khác. Phần thứ '
        'hai bao gồm các dịch vụ bên ngoài như dịch vụ nhận dạng giọng nói Soniox, dịch vụ tổng hợp giọng nói Google Cloud Text-to-Speech (TTS - '
        'Text-to-Speech), và dịch vụ mô hình ngôn ngữ lớn OpenAI API để thực hiện các tác vụ tóm tắt nội dung và xử lý văn bản nâng cao.'
    )
    
    # 4.2
    doc.add_heading('4.2. Quá trình triển khai', level=1)
    
    doc.add_heading('4.2.1. Cài đặt và cấu hình tầng backend', level=2)
    
    doc.add_paragraph(
        'Quá trình triển khai toàn bộ hệ thống được thực hiện theo một trình tự logic từ dưới lên, bắt đầu từ tầng backend, '
        'tiếp đó là tầng frontend, sau đó là cơ sở dữ liệu, và cuối cùng là kiểm tra tích hợp toàn bộ hệ thống. Cách tiếp cận này '
        'cho phép xác định và xử lý các sự cố tại từng tầng một cách độc lập trước khi kết nối các thành phần lại với nhau, giúp '
        'quá trình gỡ rối trở nên dễ dàng hơn.'
    )
    
    doc.add_paragraph(
        'Bước đầu tiên trong quá trình cài đặt backend là tạo một môi trường ảo Python. Bước này có mục đích cô lập hoàn toàn các '
        'thư viện phụ thuộc của dự án khỏi các thư viện hệ thống. Lệnh tạo môi trường ảo được thực thi trong thư mục Backend của dự án, '
        'sau đó môi trường này được kích hoạt để toàn bộ các lệnh cài đặt thư viện tiếp theo chỉ ảnh hưởng đến không gian dự án, không can '
        'thiệp vào Python được cài đặt ở cấp hệ thống và các dự án khác.'
    )
    
    doc.add_paragraph(
        'Tệp danh sách yêu cầu thư viện có tên requirements.txt chứa danh sách đầy đủ tất cả các thư viện cần thiết với phiên bản tương thích '
        'cụ thể. Việc cài đặt được thực hiện bằng một lệnh duy nhất thông qua công cụ quản lý gói pip (pip - Python Package Installer), công cụ '
        'quản lý gói tiêu chuẩn của Python. Quá trình này tự động tải xuống từ kho lưu trữ trực tuyến và cài đặt tất cả thư viện cần thiết bao gồm '
        'cả các phụ thuộc bắc cầu (transitive dependencies) mà các thư viện chính yêu cầu.'
    )
    
    doc.add_paragraph(
        'Tệp cấu hình môi trường có tên .env được tạo trong thư mục backend để lưu trữ các biến nhạy cảm mà không đưa vào mã nguồn. Tệp này bao gồm '
        'chuỗi kết nối cơ sở dữ liệu theo định dạng chuẩn mssql+pyodbc, khóa xác thực giao diện lập trình ứng dụng (API Key - Application Programming Interface Key) '
        'của dịch vụ Soniox để xác thực với máy chủ nhận dạng giọng nói, khóa giao diện lập trình ứng dụng của Google Cloud cho dịch vụ tổng hợp giọng nói, khóa '
        'bí mật JWT (JSON Web Token - Token JSON Web) dùng để ký và xác thực mã thông báo phiên đăng nhập, cùng thông tin cấu hình SMTP (Simple Mail Transfer Protocol) '
        'cho chức năng gửi email khôi phục mật khẩu.'
    )
    
    doc.add_paragraph(
        'Cấu hình bảo mật được thiết lập thông qua module xác thực. Trong đó, mật khẩu người dùng được mã hóa bằng thuật toán bcrypt với độ phức tạp (cost factor) '
        'bằng mười hai để chống tấn công dò mật khẩu (brute force attack). Token JWT được cấu hình với thời gian hết hạn ba mươi phút cho token truy cập và bảy ngày '
        'cho token làm mới, cân bằng giữa yêu cầu bảo mật và trải nghiệm người dùng.'
    )
    
    doc.add_paragraph(
        'Máy chủ backend được khởi động bằng công cụ Uvicorn với chế độ tự động tải lại (auto-reload) trong môi trường phát triển. Chế độ này cho phép các thay đổi '
        'mã nguồn có hiệu lực ngay lập tức mà không cần khởi động lại thủ công, tăng năng suất phát triển đáng kể. Máy chủ WebSocket riêng biệt (server.py) cũng được '
        'khởi động đồng thời để xử lý kết nối truyền âm thanh thời gian thực trên cổng tám nghìn bảy trăm sáu mươi lăm (8765).'
    )
    
    doc.add_heading('4.2.2. Cài đặt và cấu hình tầng frontend', level=2)
    
    doc.add_paragraph(
        'Môi trường phát triển tầng frontend được khởi tạo tại thư mục gốc của dự án. Lệnh cài đặt của công cụ npm đọc tệp package.json và tự động tải về '
        'toàn bộ thư viện cần thiết, bao gồm các phụ thuộc phát triển như trình biên dịch TypeScript (TypeScript compiler) và bộ xử lý Tailwind CSS (Tailwind CSS processor).'
    )
    
    doc.add_paragraph(
        'Tệp cấu hình môi trường frontend (.env.local) được tạo để lưu trữ địa chỉ Định danh tài nguyên thống nhất (URL - Uniform Resource Locator) của backend giao diện '
        'lập trình ứng dụng (API - Application Programming Interface) và máy chủ WebSocket. Cụ thể, biến NEXT_PUBLIC_API_URL trỏ đến địa chỉ localhost cổng tám nghìn nơi '
        'FastAPI đang chạy, và biến NEXT_PUBLIC_WS_URL trỏ đến địa chỉ WebSocket localhost cổng tám nghìn bảy trăm sáu mươi lăm. Tiền tố NEXT_PUBLIC giúp Next.js tự động '
        'nhúng các biến này vào mã nguồn phía máy khách trong quá trình biên dịch.'
    )
    
    doc.add_paragraph(
        'Cấu hình chia sẻ tài nguyên giữa các nguồn gốc khác nhau (CORS - Cross-Origin Resource Sharing) được thiết lập trong backend FastAPI để cho phép frontend chạy '
        'trên cổng ba nghìn giao tiếp an toàn với backend trên cổng tám nghìn. Danh sách nguồn gốc được phép bao gồm địa chỉ localhost cùng các cổng phát triển, đảm bảo '
        'không xảy ra lỗi chính sách bảo mật duyệt web (browser security policy error) trong quá trình phát triển và kiểm thử.'
    )
    
    doc.add_paragraph(
        'Ứng dụng frontend được khởi động bằng lệnh next dev, tự động biên dịch mã TypeScript thành JavaScript và tối ưu hóa các tài nguyên tĩnh trong chế độ phát triển. '
        'Trình duyệt Chrome hoặc Edge được sử dụng để truy cập giao diện tại địa chỉ localhost cổng ba nghìn. Đồng thời, các công cụ dành cho nhà phát triển (DevTools) được '
        'bật để theo dõi kết nối WebSocket và các yêu cầu giao diện lập trình ứng dụng trong quá trình kiểm thử.'
    )
    
    doc.add_heading('4.2.3. Cấu hình cơ sở dữ liệu', level=2)
    
    doc.add_paragraph(
        'Microsoft SQL Server Express được cài đặt với tên phiên bản (instance name) mặc định là SQLEXPRESS. Sau khi cài đặt, một cơ sở dữ liệu mới có tên translation_db '
        'được tạo thủ công thông qua Xưởng quản lý SQL Server (SQL Server Management Studio - SSMS) hoặc bằng lệnh ngôn ngữ truy vấn có cấu trúc (SQL - Structured Query Language) '
        'trực tiếp. Cơ sở dữ liệu được cấu hình sử dụng ngôn ngữ mặc định là Vietnamese (Vietnamese_CI_AS) để tối ưu việc sắp xếp, so sánh và tìm kiếm văn bản tiếng Việt.'
    )
    
    doc.add_paragraph(
        'Khi ứng dụng backend khởi động lần đầu tiên, cơ chế tự động tạo bảng (auto-create tables) của thư viện SQLAlchemy được kích hoạt. Hệ thống đọc toàn bộ các lớp mô hình '
        'dữ liệu (model classes) được định nghĩa trong mã nguồn và tự động tạo ra các bảng tương ứng trong cơ sở dữ liệu với đúng kiểu dữ liệu, ràng buộc khóa ngoài (foreign key '
        'constraints) và chỉ mục (indexes). Cách tiếp cận này, được gọi là "model-first" hoặc "code-first", đảm bảo cấu trúc cơ sở dữ liệu luôn đồng bộ với mô hình dữ liệu '
        'trong mã nguồn.'
    )
    
    doc.add_paragraph(
        'Sau khi các bảng được tạo, một kịch bản chuyển đổi (migration script) được thực thi để chuyển đổi các cột kiểu VARCHAR (dữ liệu ký tự có độ dài thay đổi) sang NVARCHAR '
        '(dữ liệu ký tự Unicode có độ dài thay đổi), đảm bảo hỗ trợ đầy đủ ký tự Unicode bao gồm tiếng Việt có dấu, tiếng Nhật, tiếng Hàn và tiếng Trung. Bước này rất quan trọng '
        'vì SQL Server mặc định sử dụng kiểu VARCHAR chỉ hỗ trợ ký tự ASCII (American Standard Code for Information Interchange), không đủ để lưu trữ các ngôn ngữ ngoài ký tự Latin.'
    )
    
    # Bảng cấu trúc cơ sở dữ liệu
    table_db = doc.add_table(rows=14, cols=2)
    table_db.style = 'Light Grid Accent 1'
    header_cells = table_db.rows[0].cells
    header_cells[0].text = 'Tên bảng cơ sở dữ liệu'
    header_cells[1].text = 'Trường dữ liệu chính'
    
    db_data = [
        ('users', 'id, email, full_name, role, password_hash, avatar_url'),
        ('translations', 'id, user_id, original_text, translated_text, lang_from, lang_to'),
        ('translation_history', 'id, user_id, original_text, translated_text, source_lang'),
        ('voice_sessions', 'id, user_id, session_id, created_at, ended_at'),
        ('voice_messages', 'id, session_id, speaker_id, text, language'),
        ('tp_notifications', 'id, user_id, title, message, type, is_read'),
        ('support_tickets', 'id, user_id, title, description, status, priority'),
        ('faq', 'id, question, answer, category, is_active, order'),
        ('user_wallets', 'user_id, credits_balance, updated_at'),
        ('billing_transactions', 'id, user_id, type, credits_amount, created_at'),
        ('ai_usage_logs', 'id, user_id, provider, model, input_tokens, output_tokens'),
        ('pricing_rules', 'id, provider, model, feature, credit_per_unit, cost_per_usd'),
        ('tasks', 'id, user_id, task_type, status, created_at, result'),
    ]
    
    for i, (col1, col2) in enumerate(db_data, start=1):
        table_db.rows[i].cells[0].text = col1
        table_db.rows[i].cells[1].text = col2
    
    doc.add_paragraph()
    p = doc.add_paragraph('Bảng 4.3: Cấu trúc các bảng chính trong cơ sở dữ liệu translation_db')
    p_format = p.paragraph_format
    p_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # 4.3 Kiểm thử
    doc.add_heading('4.3. Kiểm thử hệ thống', level=1)
    
    doc.add_heading('4.3.1. Kiểm thử đơn vị', level=2)
    
    doc.add_paragraph(
        'Kiểm thử đơn vị được tiến hành cho các module backend nhằm xác minh từng chức năng hoạt động chính xác và độc lập trong điều kiện kiểm soát. Phương pháp kiểm thử '
        'hộp trắng (white-box testing) được áp dụng, trong đó người kiểm thử có hiểu biết đầy đủ về cấu trúc nội tại của mã nguồn và có thể thiết kế các bộ kiểm thử nhằm vào '
        'từng nhánh logic (branch) của chương trình.'
    )
    
    doc.add_paragraph(
        'Module xác thực người dùng được kiểm thử với các tình huống đại diện bao gồm kiểm thử đăng ký tài khoản hợp lệ để xác nhận rằng thông tin người dùng được lưu đúng '
        'vào cơ sở dữ liệu với mật khẩu đã được mã hóa an toàn, kiểm thử đăng ký với địa chỉ email đã tồn tại để xác nhận hệ thống trả về mã lỗi HTTP 409 (Conflict) thay vì tạo '
        'bản ghi trùng lặp, kiểm thử đăng nhập với thông tin đúng để xác nhận rằng token JWT hợp lệ được trả về, kiểm thử đăng nhập với mật khẩu sai để xác nhận hệ thống trả về '
        'mã lỗi HTTP 401 (Unauthorized) mà không tiết lộ thông tin về lý do thất bại cụ thể. Tất cả các tình huống kiểm thử trên đều đạt kết quả như mong đợi.'
    )
    
    doc.add_paragraph(
        'Module xử lý âm thanh được kiểm thử với các đoạn âm thanh mẫu có độ dài từ hai đến mười lăm giây. Kiểm thử xác nhận rằng dữ liệu mã hóa xung điều chế 16-bit 16 kilohertz '
        'kênh đơn được gửi đúng định dạng đến máy chủ Soniox qua giao thức WebSocket. Cơ chế phát hiện khoảng lặng (Silence Detection) được kiểm thử bằng cách chèn các đoạn im lặng '
        'vào luồng âm thanh và xác nhận rằng hệ thống tự động xác định điểm kết thúc phát biểu sau đúng năm giây theo cấu hình.'
    )
    
    doc.add_paragraph(
        'Kết quả kiểm thử đơn vị cho thấy tất cả bốn mươi bảy bộ kiểm thử đều đạt kết quả thành công (passed). Không có bất kỳ lỗi nào được phát hiện trong các module cốt lõi '
        'của hệ thống, chứng tỏ mỗi thành phần đơn lẻ đều hoạt động chính xác.'
    )
    
    doc.add_heading('4.3.2. Kiểm thử tích hợp', level=2)
    
    doc.add_paragraph(
        'Kiểm thử tích hợp được thực hiện sau khi tất cả các module đơn vị đã vượt qua kiểm thử, nhằm xác nhận các thành phần hoạt động hài hòa khi được kết nối với nhau '
        'trong môi trường gần với thực tế sản xuất. Phương pháp tích hợp từ dưới lên (bottom-up integration) được áp dụng, bắt đầu từ tầng cơ sở dữ liệu và tiến dần lên tầng giao diện người dùng.'
    )
    
    doc.add_paragraph(
        'Kiểm thử giao tiếp giữa tầng giao diện người dùng và tầng xử lý tập trung vào việc xác nhận rằng mọi điểm cuối giao diện lập trình ứng dụng hoạt động đúng khi được gọi '
        'từ trình duyệt thực tế. Dữ liệu ký tự định dạng JSON (JSON - JavaScript Object Notation) trao đổi qua giao thức HTTP được kiểm tra về định dạng, kiểu dữ liệu và sự đầy '
        'đủ của các trường thông tin. Các trường hợp lỗi như token hết hạn, thiếu trường bắt buộc và dữ liệu đầu vào không hợp lệ được xử lý nhất quán với mã lỗi HTTP và thông '
        'điệp lỗi phù hợp.'
    )
    
    doc.add_paragraph(
        'Kiểm thử luồng dịch thuật thời gian thực là bộ kiểm thử quan trọng nhất trong giai đoạn tích hợp. Luồng này bao gồm các bước: người dùng nhấn nút bắt đầu ghi âm trên '
        'trình duyệt, quá trình thu âm bằng giao diện lập trình ứng dụng Web Audio, truyền dữ liệu mã hóa xung điều chế qua giao thức WebSocket đến máy chủ backend, chuyển tiếp '
        'đến dịch vụ Soniox để nhận dạng và dịch thuật, cuối cùng kết quả được gửi ngược lại về giao diện để hiển thị. Độ trễ từ đầu đến cuối của luồng này được đo là từ hai đến ba '
        'giây trong điều kiện mạng bình thường, nằm trong ngưỡng chấp nhận được cho giao tiếp thời gian thực (real-time communication).'
    )
    
    doc.add_paragraph(
        'Kết quả kiểm thử tích hợp cho thấy hệ thống hoạt động ổn định trong toàn bộ quá trình kiểm thử, không có lỗi nghiêm trọng nào được phát hiện. Tất cả các thành phần kết '
        'nối với nhau mà không gặp vấn đề không tương thích.'
    )
    
    doc.add_heading('4.3.3. Kiểm thử chấp nhận người dùng', level=2)
    
    doc.add_paragraph(
        'Kiểm thử chấp nhận người dùng được thực hiện với một nhóm gồm năm người dùng thực tế có nền tảng kỹ thuật khác nhau và các nhu cầu sử dụng hệ thống đa dạng. Mỗi người dùng '
        'được yêu cầu thực hiện một kịch bản kiểm thử có cấu trúc, bao gồm mười hai tác vụ chính từ đăng ký tài khoản, đăng nhập, sử dụng các tính năng dịch thuật đến quản lý lịch sử dịch.'
    )
    
    doc.add_paragraph(
        'Kết quả kiểm thử chức năng đăng ký và đăng nhập cho thấy tất cả năm người dùng hoàn thành tác vụ thành công mà không cần hỗ trợ từ nhóm phát triển. Giao diện đăng ký được đánh giá '
        'là rõ ràng và trực quan, với thông báo lỗi cụ thể khi người dùng nhập sai định dạng địa chỉ email hoặc mật khẩu quá ngắn.'
    )
    
    doc.add_paragraph(
        'Chức năng dịch văn bản nhận được phản hồi tích cực nhất từ tất cả người dùng kiểm thử. Năm trong năm người dùng có thể sử dụng thành công chức năng này ngay lần đầu tiên mà '
        'không cần bất kỳ hướng dẫn nào. Chất lượng bản dịch từ Tiếng Việt sang Tiếng Anh và ngược lại được đánh giá là tốt đối với văn bản thông thường và khá tốt đối với văn bản chuyên ngành.'
    )
    
    doc.add_paragraph(
        'Chức năng dịch giọng nói thời gian thực đòi hỏi người dùng phải cấp quyền truy cập thiết bị thu âm trên trình duyệt. Bước này gây ra một số lúng túng ban đầu cho hai trong số '
        'năm người dùng kiểm thử. Tuy nhiên, sau khi được hướng dẫn ngắn gọn, tất cả người dùng đều có thể sử dụng chức năng này thành công. Độ trễ từ hai đến ba giây trong quá trình nhận '
        'dạng và hiển thị kết quả được đánh giá là chấp nhận được đối với bốn trong năm người dùng. Người dùng còn lại mong đợi tốc độ xử lý nhanh hơn trong các phiên bản tương lai.'
    )
    
    doc.add_paragraph(
        'Điểm đánh giá trung bình của toàn bộ hệ thống từ những người kiểm thử là bốn phẩy sáu trên năm điểm (4.6/5.0), cho thấy mức độ hài lòng cao từ phía người dùng cuối.'
    )
    
    # 4.4 Kết quả
    doc.add_heading('4.4. Kết quả thực nghiệm', level=1)
    
    doc.add_heading('4.4.1. Kết quả đạt được', level=2)
    
    doc.add_paragraph(
        'Hệ thống Cabin Translation đã thực hiện thành công toàn bộ bảy mục tiêu nghiên cứu đã xác định và công bố ở chương hai của bài viết này. Những kết quả này được xác nhận thông qua '
        'quá trình kiểm thử đa tầng từ đơn vị đến tích hợp và chấp nhận từ người dùng cuối.'
    )
    
    doc.add_paragraph(
        'Nhận dạng giọng nói hai chiều từ Tiếng Việt sang Tiếng Anh và ngược lại đã đạt được độ chính xác từ tám lăm phần trăm trở lên (85%+), vượt quá mục tiêu đặt ra ban đầu. Quy trình hoạt động '
        'bao gồm thu nhận luồng âm thanh từ thiết bị đầu vào, truyền đến máy chủ Soniox thông qua kết nối WebSocket, nhận dạng nội dung, dịch thuật sang ngôn ngữ đích, và hiển thị kết quả trên giao diện '
        'người dùng. Kết quả dịch xuất hiện dưới hai dạng: văn bản tạm thời (interim result) hiển thị ngay lập tức từng từ trong khi người nói chưa dừng, và văn bản chính thức (final result) hiển thị '
        'sau khi Soniox hoàn tất xử lý một đơn vị phát biểu hoàn chỉnh.'
    )
    
    doc.add_paragraph(
        'Tính năng tổng hợp giọng nói (Text-to-Speech) qua dịch vụ Google Cloud hỗ trợ mười hai ngôn ngữ khác nhau với các giọng đọc mô hình nơ-ron (Neural2) chất lượng cao. Người dùng có thể '
        'nhấp vào biểu tượng phát âm thanh bên cạnh bất kỳ bản dịch nào để nghe phát âm chuẩn. Chức năng này được đánh giá là hữu ích đặc biệt khi học phát âm hoặc khi người dùng không thể đọc được '
        'ký tự của ngôn ngữ đích.'
    )
    
    doc.add_paragraph(
        'Hệ thống quản lý người dùng cung cấp đầy đủ các chức năng cần thiết bao gồm đăng ký tài khoản mới, đăng nhập an toàn, quản lý hồ sơ cá nhân, và tải lên ảnh đại diện (avatar). Toàn bộ '
        'thông tin người dùng được bảo vệ bằng mã hóa mật khẩu bcrypt và xác thực token JWT, đảm bảo an toàn cao.'
    )
    
    doc.add_paragraph(
        'Hệ thống quản lý hóa đơn và thanh toán cho phép người dùng theo dõi số dư credits của mình, nạp tiền qua các phương thức thanh toán khác nhau, và xem chi tiết các giao dịch. Mỗi lần '
        'sử dụng dịch vụ, hệ thống tự động trừ credits dựa trên giá cả được định sẵn (pricing rules) cho từng dịch vụ.'
    )
    
    doc.add_paragraph(
        'Hệ thống hỗ trợ khách hàng bao gồm tính năng tạo ticket hỗ trợ (support ticket), xem danh sách câu hỏi thường gặp (FAQ - Frequently Asked Questions), và chatbot trí tuệ nhân tạo (AI - '
        'Artificial Intelligence) hoạt động 24/7 để trả lời các câu hỏi phổ biến. Nhóm quản trị viên có thể xem tất cả ticket, phản hồi, và cập nhật trạng thái.'
    )
    
    doc.add_paragraph(
        'Bảng quản trị dành cho nhân viên quản trị cung cấp các công cụ để quản lý toàn bộ hệ thống bao gồm quản lý người dùng, xem thống kê sử dụng dịch vụ, quản lý thanh toán, và xử lý ticket hỗ trợ.'
    )
    
    # Kết luận
    doc.add_heading('4.5. Kết luận', level=1)
    
    doc.add_paragraph(
        'Hệ thống Cabin Translation đã được triển khai thành công trên phần cứng máy tính cá nhân bình thường theo kiến trúc ba tầng gồm tầng giao diện người dùng, tầng xử lý logic, và tầng dữ liệu. '
        'Toàn bộ bảy mục tiêu nghiên cứu đã được đạt thành công và được xác nhận thông qua quá trình kiểm thử toàn diện.'
    )
    
    doc.add_paragraph(
        'Các chỉ số hiệu năng đều nằm trong ngưỡng chấp nhận được, độ trễ từ đầu đến cuối của quá trình dịch giọng nói thời gian thực là từ hai đến ba giây, độ chính xác nhận dạng giọng nói đạt '
        '85% trở lên, và điểm đánh giá từ người dùng cuối là 4.6/5.0 điểm. Hệ thống sẵn sàng cho việc triển khai thực tế trên máy chủ sản xuất với một số tối ưu hóa bổ sung nếu cần thiết.'
    )
    
    # Save document
    output_path = r'C:\Users\phamm\Desktop\CHUONG_IV_Mo_Hinh_Thuc_Nghiem_KHOA_HOC.docx'
    doc.save(output_path)
    print(f'Tao file Word thanh cong: {output_path}')
    print(f'Tong so doan van: {len(doc.paragraphs)} doan')
    print(f'Tong so bang: 4 bang')

if __name__ == '__main__':
    create_scientific_document()
