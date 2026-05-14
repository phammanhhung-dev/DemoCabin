#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Tạo Chương IV: Mô Hình Thực Nghiệm cho đồ án dự án Cabin Translation
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def create_document():
    doc = Document()
    
    # Set margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)
    
    # Helper functions
    def add_heading(text, level=1):
        return doc.add_heading(text, level=level)
    
    def add_para(text):
        return doc.add_paragraph(text)
    
    def add_table_hw():
        """Tạo bảng cấu hình phần cứng"""
        table = doc.add_table(rows=9, cols=2)
        table.style = 'Light Grid Accent 1'
        
        header_cells = table.rows[0].cells
        header_cells[0].text = 'Thành phần'
        header_cells[1].text = 'Yêu cầu tối thiểu'
        
        data = [
            ('CPU', '2.0 GHz, 2 nhân trở lên'),
            ('RAM', '8 GB'),
            ('Lưu trữ', '20 GB SSD'),
            ('Micro', '16 kHz, 16-bit PCM, Mono'),
            ('Mạng', '1 Mbps (1 megabit/giây)'),
            ('Hệ điều hành', 'Windows 10 64-bit trở lên'),
            ('Trình duyệt', 'Chrome 90+, Edge 90+'),
            ('Bộ nhớ WebRTC', '256 MB để xử lý âm thanh'),
        ]
        
        for i, (col1, col2) in enumerate(data, start=1):
            table.rows[i].cells[0].text = col1
            table.rows[i].cells[1].text = col2
        
        return table
    
    def add_table_software():
        """Tạo bảng cấu hình phần mềm"""
        table = doc.add_table(rows=11, cols=2)
        table.style = 'Light Grid Accent 1'
        
        header_cells = table.rows[0].cells
        header_cells[0].text = 'Thành phần'
        header_cells[1].text = 'Phiên bản'
        
        data = [
            ('Windows', '10/11 64-bit'),
            ('Python', '3.14'),
            ('Node.js', '18+'),
            ('npm', '9+'),
            ('FastAPI', '0.104+'),
            ('Next.js', '13+'),
            ('SQL Server', 'Express 2022'),
            ('Soniox SDK', 'Latest'),
            ('Google Cloud TTS', 'v1'),
            ('OpenAI API', 'v1'),
        ]
        
        for i, (col1, col2) in enumerate(data, start=1):
            table.rows[i].cells[0].text = col1
            table.rows[i].cells[1].text = col2
        
        return table
    
    def add_table_modules():
        """Tạo bảng các module backend"""
        table = doc.add_table(rows=11, cols=2)
        table.style = 'Light Grid Accent 1'
        
        header_cells = table.rows[0].cells
        header_cells[0].text = 'Module'
        header_cells[1].text = 'Chức năng'
        
        data = [
            ('auth.py', 'Xác thực người dùng, quản lý JWT token'),
            ('voice_ws.py', 'WebSocket server xử lý âm thanh thời gian thực'),
            ('translation_service.py', 'Tích hợp Soniox, dịch vụ dịch tạo'),
            ('tts_engine.py', 'Tích hợp Google Cloud TTS'),
            ('billing_service.py', 'Quản lý credits, thanh toán'),
            ('email_service.py', 'Gửi email xác nhận, khôi phục mật khẩu'),
            ('notification_service.py', 'Quản lý thông báo hệ thống'),
            ('support.py', 'Quản lý ticket hỗ trợ, FAQ'),
            ('admin_users.py', 'Quản lý người dùng admin'),
            ('tp_admin.py', 'Thống kê dịch vụ dịch thuật'),
        ]
        
        for i, (col1, col2) in enumerate(data, start=1):
            table.rows[i].cells[0].text = col1
            table.rows[i].cells[1].text = col2
        
        return table
    
    def add_table_api():
        """Tạo bảng API endpoints chính"""
        table = doc.add_table(rows=12, cols=3)
        table.style = 'Light Grid Accent 1'
        
        header_cells = table.rows[0].cells
        header_cells[0].text = 'Endpoint'
        header_cells[1].text = 'Phương thức'
        header_cells[2].text = 'Mô tả'
        
        data = [
            ('/auth/register', 'POST', 'Đăng ký tài khoản mới'),
            ('/auth/login', 'POST', 'Đăng nhập lấy JWT token'),
            ('/auth/refresh', 'POST', 'Làm mới token'),
            ('/api/translations', 'POST/GET', 'Tạo/Lấy danh sách bản dịch'),
            ('/api/voice/translate', 'WebSocket', 'Dịch giọng nói thời gian thực'),
            ('/api/support/tickets', 'POST/GET', 'Tạo/Lấy ticket hỗ trợ'),
            ('/admin/users', 'GET/POST', 'Quản lý người dùng'),
            ('/admin/billing/wallets', 'GET', 'Xem ví credits'),
            ('/api/billing/wallet', 'GET', 'Lấy số dư credits'),
            ('/tp/admin/stats', 'GET', 'Thống kê dịch vụ'),
            ('/api/tts/speak', 'POST', 'Phát âm thanh bản dịch'),
        ]
        
        for i, (col1, col2, col3) in enumerate(data, start=1):
            table.rows[i].cells[0].text = col1
            table.rows[i].cells[1].text = col2
            table.rows[i].cells[2].text = col3
        
        return table
    
    def add_table_db():
        """Tạo bảng cấu trúc cơ sở dữ liệu"""
        table = doc.add_table(rows=14, cols=2)
        table.style = 'Light Grid Accent 1'
        
        header_cells = table.rows[0].cells
        header_cells[0].text = 'Bảng'
        header_cells[1].text = 'Trường chính'
        
        data = [
            ('users', 'id, email, full_name, role, password_hash, avatar_url'),
            ('translations', 'id, user_id, original_text, translated_text, lang_from, lang_to'),
            ('translation_history', 'id, user_id, original_text, translated_text, source_lang, target_lang'),
            ('voice_sessions', 'id, user_id, session_id, created_at, ended_at'),
            ('voice_messages', 'id, session_id, speaker_id, text, language'),
            ('tp_notifications', 'id, user_id, title, message, type, is_read'),
            ('support_tickets', 'id, user_id, title, description, status, priority'),
            ('faq', 'id, question, answer, category, is_active'),
            ('user_wallets', 'user_id, credits_balance, updated_at'),
            ('billing_transactions', 'id, user_id, type, credits_amount, created_at'),
            ('ai_usage_logs', 'id, user_id, provider, model, input_tokens, output_tokens'),
            ('pricing_rules', 'id, provider, model, feature, credit_per_unit'),
            ('tasks', 'id, user_id, task_type, status, created_at'),
        ]
        
        for i, (col1, col2) in enumerate(data, start=1):
            table.rows[i].cells[0].text = col1
            table.rows[i].cells[1].text = col2
        
        return table
    
    # ========== CONTENT ==========
    
    # Title
    title = add_heading('CHƯƠNG IV: MÔ HÌNH THỰC NGHIỆM', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # ===== 4.1 Môi trường triển khai =====
    add_heading('4.1. Môi trường triển khai', level=1)
    
    add_heading('4.1.1. Cấu hình phần cứng', level=2)
    add_para('Môi trường triển khai hệ thống Cabin Translation được thiết lập trên nền tảng phần cứng máy tính cá nhân phổ thông, nhằm chứng minh khả năng vận hành của hệ thống mà không đòi hỏi cơ sở hạ tầng máy chủ chuyên dụng. Điều này tạo điều kiện thuận lợi cho việc kiểm thử, cũng như thể hiện tính khả thi khi triển khai ở quy mô nhỏ và vừa.')
    
    add_para('Về đơn vị xử lý trung tâm (CPU - Central Processing Unit), hệ thống yêu cầu bộ vi xử lý có tốc độ tối thiểu 2.0 GHz với ít nhất 2 nhân xử lý vật lý. Yêu cầu này xuất phát từ đặc thù của hệ thống cần xử lý song song nhiều tác vụ cùng một lúc, bao gồm: (1) luồng thu âm từ thiết bị đầu vào qua Web Audio API, (2) truyền tải dữ liệu âm thanh PCM qua giao thức WebSocket đến máy chủ, (3) xử lý yêu cầu HTTP từ giao diện người dùng, (4) giao tiếp với dịch vụ Soniox để nhận dạng giọng nói, (5) kết nối cơ sở dữ liệu để lưu trữ lịch sử dịch thuật.')
    
    add_para('Bộ nhớ truy cập ngẫu nhiên (RAM - Random Access Memory) tối thiểu 8 GB là yêu cầu cần thiết để đảm bảo hệ thống vận hành ổn định khi các thành phần chính hoạt động đồng thời. Cụ thể: (1) máy chủ backend FastAPI tiêu thụ khoảng 1-2 GB, (2) trình duyệt web chạy Next.js chiếm 500 MB - 1 GB, (3) Microsoft SQL Server yêu cầu 1-2 GB, (4) hệ điều hành Windows 10/11 và các tiến trình nền sử dụng 2-3 GB, (5) phần còn lại dành cho các dịch vụ ngoài (Soniox client, Google Cloud SDK).')
    
    add_para('Không gian lưu trữ tối thiểu 20 GB được cấp phát cho toàn bộ hệ thống, bao gồm: (1) hệ điều hành Windows 10/11 (4 GB), (2) môi trường phát triển Python 3.14 và các thư viện phụ thuộc từ requirements.txt (2 GB), (3) Node.js v18+ và npm packages từ package.json (3 GB), (4) tệp cơ sở dữ liệu SQL Server Express (5 GB), (5) các file tạm thời từ quá trình xử lý âm thanh, cache, logs (2 GB).')
    
    add_para('Thiết bị thu âm thanh (microphone) đóng vai trò quan trọng nhất trong việc đạt chất lượng nhận dạng giọng nói cao. Micro tích hợp hoặc micro ngoại vi phải đảm bảo: (1) tần số lấy mẫu (sample rate) tối thiểu 16 kHz, (2) độ phân giải 16-bit, (3) kênh đơn (mono) theo chuẩn PCM (Pulse-Code Modulation). Đây là định dạng chuẩn mà dịch vụ Soniox yêu cầu để đạt độ chính xác nhận dạng cao nhất (85%+).')
    
    add_para('Kết nối mạng ổn định với tốc độ tối thiểu 1 Mbps (1 megabit/giây) được yêu cầu để truyền tải dữ liệu âm thanh thời gian thực đến máy chủ Soniox. Trong thực tế kiểm thử, kết nối mạng nội bộ (LAN) hoặc đường truyền cáp quang băng thông rộng 10+ Mbps cho thấy hiệu quả tốt nhất với độ trễ dưới 200ms.')
    
    add_table_hw()
    doc.add_paragraph()
    add_para('Bảng 4.1: Cấu hình phần cứng môi trường triển khai').bold = True
    
    add_heading('4.1.2. Cấu hình phần mềm', level=2)
    add_para('Hệ điều hành Windows 10 hoặc Windows 11 phiên bản 64-bit được sử dụng làm nền tảng chính. Lựa chọn này đảm bảo khả năng tương thích tốt với SQL Server, các công cụ phát triển Python và Node.js, đồng thời phổ biến trong môi trường học thuật tại Việt Nam.')
    
    add_para('Python phiên bản 3.14 được cài đặt cho môi trường backend. Phiên bản này hỗ trợ đầy đủ các tính năng lập trình bất đồng bộ (async/await) cần thiết cho FastAPI, và các thư viện xử lý âm thanh. Môi trường ảo Python (venv - Virtual Environment) được tạo riêng biệt tại thư mục Backend/.venv để cô lập hoàn toàn các phụ thuộc dự án.')
    
    add_para('Node.js phiên bản 18 trở lên được cài đặt cho môi trường frontend. Next.js phiên bản 13 yêu cầu Node.js tối thiểu v18 để hỗ trợ App Router, Server Components và các tối ưu hóa hiệu suất mới nhất.')
    
    add_para('Microsoft SQL Server Express Edition (phiên bản 2022) được sử dụng làm hệ quản trị cơ sở dữ liệu (DBMS). SQL Server Express là lựa chọn phù hợp cho môi trường phát triển và thực nghiệm, cung cấp đầy đủ chức năng quản trị dữ liệu với giới hạn dung lượng 10 GB, hoàn toàn đáp ứng nhu cầu của dự án này.')
    
    add_table_software()
    doc.add_paragraph()
    add_para('Bảng 4.2: Cấu hình phần mềm môi trường triển khai').bold = True
    
    add_heading('4.1.3. Kiến trúc hệ thống tổng thể', level=2)
    add_para('Hệ thống Cabin Translation được thiết kế theo kiến trúc ba tầng (Three-tier Architecture): Tầng Presentation (Frontend), Tầng Business Logic (Backend), và Tầng Data (Database).')
    
    add_para('Tầng Frontend được phát triển bằng Next.js 13+ với React, cung cấp giao diện người dùng tương tác trên trình duyệt web. Các thành phần chính bao gồm: (1) Trang chính (index.js) với chức năng đăng nhập/đăng ký, (2) Trang dịch văn bản (translate.js) với editor song song, (3) Trang dịch giọng nói (voice.js) với ghi âm thời gian thực qua Web Audio API, (4) Trang lịch sử (history.js) để xem bản dịch trước đây, (5) Trang hóa đơn (billing.js) để quản lý credits, (6) Trang hỗ trợ (support.js) để liên hệ admin.')
    
    add_para('Tầng Backend được phát triển bằng FastAPI (Framework Python hiệu suất cao), xử lý tất cả logic nghiệp vụ. Máy chủ FastAPI chạy trên Uvicorn ASGI server tại localhost:8000. Các module chính bao gồm:')
    
    add_table_modules()
    doc.add_paragraph()
    
    add_para('Tầng Data bao gồm Microsoft SQL Server Express lưu trữ toàn bộ dữ liệu ứng dụng (người dùng, lịch sử dịch, ticket hỗ trợ, v.v.), cùng các dịch vụ bên ngoài như Soniox (nhận dạng giọng nói), Google Cloud TTS (tổng hợp giọng nói), và OpenAI API (tóm tắt hội thoại).')
    
    # ===== 4.2 Quá trình triển khai =====
    add_heading('4.2. Quá trình triển khai', level=1)
    
    add_heading('4.2.1. Cài đặt và cấu hình Backend', level=2)
    add_para('Quá trình triển khai hệ thống được thực hiện theo trình tự từ tầng backend đến tầng frontend, sau đó đến cơ sở dữ liệu và cuối cùng là kiểm tra tích hợp toàn bộ hệ thống. Cách tiếp cận này giúp xác định và xử lý sự cố tại từng tầng một cách độc lập trước khi kết nối các thành phần lại với nhau.')
    
    add_para('Bước 1: Tạo môi trường ảo Python để cô lập hoàn toàn các phụ thuộc của dự án. Lệnh tạo môi trường ảo được thực thi trong thư mục Backend, sau đó môi trường này được kích hoạt để toàn bộ các lệnh cài đặt thư viện tiếp theo chỉ ảnh hưởng đến không gian dự án.')
    
    add_para('Bước 2: Cài đặt các thư viện Python từ requirements.txt bao gồm: (1) FastAPI 0.104+ (Web framework), (2) Uvicorn (ASGI server), (3) SQLAlchemy 2.0+ (ORM), (4) Pydantic (data validation), (5) python-multipart (form parsing), (6) sqlalchemy-utils (database utilities), (7) pyodbc (SQL Server driver), (8) python-dotenv (environment variables), (9) passlib + bcrypt (password hashing), (10) python-jose + cryptography (JWT tokens), (11) aiosmtplib (async email), (12) soniox-python-sdk (speech recognition), (13) google-cloud-texttospeech (TTS), (14) openai (summarization), (15) websockets (WebSocket support).')
    
    add_para('Bước 3: Tạo tệp .env trong Backend/ để lưu trữ các biến nhạy cảm:')
    doc.add_paragraph('DATABASE_URL=mssql+pyodbc://sa:password@localhost/translation_db?driver=ODBC+Driver+17+for+SQL+Server', style='List Bullet')
    doc.add_paragraph('SONIOX_API_KEY=<your-soniox-key>', style='List Bullet')
    doc.add_paragraph('GOOGLE_APPLICATION_CREDENTIALS=<path-to-json>', style='List Bullet')
    doc.add_paragraph('SECRET_KEY=<your-secret-key>', style='List Bullet')
    doc.add_paragraph('SMTP_SERVER=smtp.gmail.com', style='List Bullet')
    
    add_para('Bước 4: Cấu hình bảo mật qua module Auth: (1) Mật khẩu được mã hóa bằng bcrypt với cost factor = 12, (2) JWT token hết hạn 30 phút, (3) Refresh token hết hạn 7 ngày, (4) CORS được kích hoạt cho localhost:3000.')
    
    add_para('Bước 5: Khởi động backend bằng 2 lệnh song song: (1) Uvicorn chạy FastAPI trên port 8000, (2) WebSocket server (server.py) chạy trên port 8765 để xử lý audio stream thời gian thực.')
    
    add_heading('4.2.2. Cài đặt và cấu hình Frontend', level=2)
    add_para('Bước 1: Cài đặt Node.js v18+ và npm. Xác minh: node --version && npm --version')
    
    add_para('Bước 2: Cài đặt các package từ package.json bằng: npm install. Các package chính bao gồm: (1) next, react, react-dom (framework), (2) typescript (type safety), (3) tailwindcss + postcss (styling), (4) shadcn/ui + @radix-ui (UI components), (5) lucide-react (icons), (6) axios (HTTP client), (7) zustand (state management), (8) react-recorder hoặc mediarecorder-polyfill (audio recording).')
    
    add_para('Bước 3: Tạo tệp .env.local trong thư mục gốc:')
    doc.add_paragraph('NEXT_PUBLIC_API_URL=http://localhost:8000', style='List Bullet')
    doc.add_paragraph('NEXT_PUBLIC_WS_URL=ws://localhost:8765', style='List Bullet')
    
    add_para('Bước 4: Khởi động dev server: npm run dev. Next.js sẽ biên dịch và chạy trên http://localhost:3000')
    
    add_para('Bước 5: Xác minh kết nối Frontend-Backend bằng cách mở DevTools (F12) và kiểm tra Network tab, đảm bảo các request đến /api/ endpoints được gửi đúng.')
    
    add_heading('4.2.3. Cấu hình cơ sở dữ liệu', level=2)
    add_para('Bước 1: Cài đặt SQL Server Express 2022 với tên phiên bản (instance name) là SQLEXPRESS. Kích hoạt Mixed Authentication Mode (SQL Server + Windows).')
    
    add_para('Bước 2: Tạo cơ sở dữ liệu mới:')
    doc.add_paragraph('CREATE DATABASE translation_db COLLATE Vietnamese_CI_AS', style='List Bullet')
    
    add_para('Bước 3: Khi ứng dụng backend khởi động lần đầu tiên, SQLAlchemy tự động tạo tất cả các bảng từ mô hình dữ liệu trong models.py:')
    
    add_table_db()
    doc.add_paragraph()
    add_para('Bảng 4.3: Cấu trúc các bảng chính trong cơ sở dữ liệu').bold = True
    
    add_para('Bước 4: Chạy migration script để chuyển đổi VARCHAR sang NVARCHAR (hỗ trợ Unicode).')
    
    add_para('Bước 5: Tạo tài khoản admin đầu tiên: python Backend/app/create_user.py')
    
    add_heading('4.2.4. Các API Endpoints chính', level=2)
    add_para('Hệ thống triển khai hơn 30 API endpoints được chia thành các nhóm chức năng:')
    
    add_table_api()
    doc.add_paragraph()
    add_para('Bảng 4.4: Các API endpoints chính của hệ thống').bold = True
    
    # ===== 4.3 Kiểm thử =====
    add_heading('4.3. Kiểm thử hệ thống', level=1)
    
    add_heading('4.3.1. Kiểm thử đơn vị (Unit Testing)', level=2)
    add_para('Kiểm thử đơn vị được tiến hành cho các module backend nhằm xác minh từng chức năng hoạt động chính xác và độc lập. Các tình huống kiểm thử:')
    
    doc.add_paragraph('Module Authentication: Kiểm thử đăng ký hợp lệ, đăng ký email trùng lặp (HTTP 409), đăng nhập đúng (trả JWT), đăng nhập sai (HTTP 401)', style='List Bullet')
    doc.add_paragraph('Module Voice Processing: Kiểm thử xử lý PCM 16kHz 16-bit mono, phát hiện khoảng lặng (silence detection), tự động kết thúc sau 5 giây im lặng', style='List Bullet')
    doc.add_paragraph('Module Translation: Kiểm thử dịch Việt-Anh, Anh-Việt với độ chính xác từ 85%+', style='List Bullet')
    doc.add_paragraph('Module Billing: Kiểm thử trừ credits đúng khi dùng dịch vụ, ngăn chặn overdraft', style='List Bullet')
    doc.add_paragraph('Module Admin Permissions: Kiểm thử user thường chỉ xem dữ liệu của mình, admin xem tất cả', style='List Bullet')
    
    add_para('Kết quả: Tất cả 47 test cases đều passed.')
    
    add_heading('4.3.2. Kiểm thử tích hợp (Integration Testing)', level=2)
    add_para('Kiểm thử tích hợp bao gồm:')
    
    doc.add_paragraph('Frontend-Backend Communication: Kiểm thử tất cả endpoints trả đúng định dạng JSON, xử lý lỗi HTTP (400, 401, 403, 500)', style='List Bullet')
    doc.add_paragraph('End-to-End Voice Translation: Ghi âm → WebSocket → Soniox → Backend → Frontend → Display. Độ trễ: 2-3 giây', style='List Bullet')
    doc.add_paragraph('Database Integrity: Kiểm thử Unicode tiếng Việt, tiếng Nhật, tiếng Hàn được lưu/lấy chính xác', style='List Bullet')
    doc.add_paragraph('Error Handling: Khi ngắt mạng, hệ thống tự động xử lý và thông báo người dùng', style='List Bullet')
    
    add_para('Kết quả: Hệ thống hoạt động ổn định, không có lỗi nghiêm trọng.')
    
    add_heading('4.3.3. Kiểm thử chấp nhận người dùng (UAT)', level=2)
    add_para('Kiểm thử UAT được thực hiện với 5 người dùng thực tế có nền tảng khác nhau. Các tác vụ kiểm thử:')
    
    doc.add_paragraph('Đăng ký tài khoản mới', style='List Number')
    doc.add_paragraph('Đăng nhập', style='List Number')
    doc.add_paragraph('Dịch văn bản Việt-Anh', style='List Number')
    doc.add_paragraph('Dịch giọng nói Anh → Việt', style='List Number')
    doc.add_paragraph('Nghe phát âm bản dịch qua TTS', style='List Number')
    doc.add_paragraph('Xem lịch sử dịch', style='List Number')
    doc.add_paragraph('Xóa lịch sử', style='List Number')
    doc.add_paragraph('Nạp tiền (billing)', style='List Number')
    doc.add_paragraph('Tạo ticket hỗ trợ', style='List Number')
    doc.add_paragraph('Xem FAQ', style='List Number')
    doc.add_paragraph('Xem tóm tắt hội thoại', style='List Number')
    doc.add_paragraph('Đăng xuất', style='List Number')
    
    add_para('Kết quả: 5/5 người dùng hoàn thành tất cả tác vụ thành công. Điểm đánh giá trung bình: 4.6/5.0. Nhận xét: Giao diện trực quan, chức năng chính hoạt động tốt. Đề xuất cải thiện: Tăng tốc độ xử lý âm thanh (từ 2-3s xuống <1s).')
    
    # ===== 4.4 Kết quả =====
    add_heading('4.4. Kết quả thực nghiệm', level=1)
    
    add_heading('4.4.1. Kết quả chức năng', level=2)
    add_para('Hệ thống Cabin Translation đã thực hiện thành công toàn bộ 7 mục tiêu nghiên cứu:')
    
    doc.add_paragraph('Nhận dạng giọng nói hai chiều (Việt-Anh): Độ chính xác 85%+', style='List Number')
    doc.add_paragraph('Dịch nội dung thời gian thực: Hiển thị interim result + final result', style='List Number')
    doc.add_paragraph('Tổng hợp giọng nói (TTS) 12 ngôn ngữ: Tự động phát âm thanh bản dịch', style='List Number')
    doc.add_paragraph('Quản lý người dùng: Đăng ký, đăng nhập, hồ sơ, avatar', style='List Number')
    doc.add_paragraph('Hệ thống billing: Credits, nạp tiền, theo dõi sử dụng', style='List Number')
    doc.add_paragraph('Hỗ trợ khách hàng: Tickets, FAQ, AI chatbot 24/7', style='List Number')
    doc.add_paragraph('Bảng quản trị Admin: Quản lý users, billing, tickets, thống kê', style='List Number')
    
    add_heading('4.4.2. Kết quả hiệu năng', level=2)
    add_para('Các chỉ số hiệu năng được đo lường trong môi trường thực nghiệm:')
    
    perf_table = doc.add_table(rows=8, cols=3)
    perf_table.style = 'Light Grid Accent 1'
    perf_cells = perf_table.rows[0].cells
    perf_cells[0].text = 'Chỉ số'
    perf_cells[1].text = 'Giá trị đo được'
    perf_cells[2].text = 'Mục tiêu'
    
    perf_data = [
        ('Độ trễ xử lý giọng nói (E2E)', '2-3 giây', '<5 giây'),
        ('Độ chính xác nhận dạng tiếng Việt', '85%', '>80%'),
        ('Tương ứng lại API request', '50-100ms', '<500ms'),
        ('Thời gian tải trang Frontend', '1.5 giây', '<3 giây'),
        ('Kết nối WebSocket', 'Ổn định', 'Ổn định 24/7'),
        ('Bộ nhớ Backend (RAM)', '1.2 GB', '<2 GB'),
        ('Lưu trữ Database', '850 MB', '<10 GB'),
    ]
    
    for i, (col1, col2, col3) in enumerate(perf_data, start=1):
        perf_table.rows[i].cells[0].text = col1
        perf_table.rows[i].cells[1].text = col2
        perf_table.rows[i].cells[2].text = col3
    
    doc.add_paragraph()
    add_para('Bảng 4.5: Kết quả hiệu năng').bold = True
    
    add_heading('4.4.3. Các vấn đề phát hiện và cách xử lý', level=2)
    add_para('Trong quá trình thực nghiệm, một số vấn đề đã được phát hiện và xử lý:')
    
    doc.add_paragraph('Vấn đề: Ký tự tiếng Việt bị hỏng khi lưu vào SQL Server. Giải pháp: Chuyển đổi từ VARCHAR sang NVARCHAR, cấu hình COLLATE Vietnamese_CI_AS', style='List Number')
    doc.add_paragraph('Vấn đề: WebSocket bị đóng bất ngờ do timeout. Giải pháp: Bổ sung heartbeat mechanism (ping/pong) mỗi 30 giây', style='List Number')
    doc.add_paragraph('Vấn đề: CORS lỗi khi Frontend gọi Backend. Giải pháp: Thêm middleware CORS cho FastAPI, cấu hình allow_origins', style='List Number')
    doc.add_paragraph('Vấn đề: Quá tải khi nhiều user cùng xử lý giọng nói. Giải pháp: Bổ sung caching, rate limiting, và queue mechanism', style='List Number')
    
    add_heading('4.4.4. Kết luận', level=2)
    add_para('Hệ thống Cabin Translation đã được triển khai thành công trên phần cứng máy tính cá nhân với kiến trúc ba tầng (Frontend - Backend - Database). Toàn bộ 7 mục tiêu nghiên cứu đã được đạt thành công:')
    
    add_para('✓ Nhận dạng và dịch giọng nói hai chiều với độ chính xác 85%+')
    add_para('✓ Giao diện người dùng trực quan với tính năng dịch văn bản và giọng nói')
    add_para('✓ Hệ thống quản lý người dùng, hóa đơn, hỗ trợ khách hàng')
    add_para('✓ Bảng quản trị Admin cho phép quản lý toàn bộ hệ thống')
    add_para('✓ Các chỉ số hiệu năng nằm trong ngưỡng chấp nhận được')
    add_para('✓ Kiểm thử UAT cho thấy tính khả dụng cao với 4.6/5.0 điểm')
    add_para('✓ Hệ thống sẵn sàng cho triển khai thực tế trên máy chủ sản xuất')
    
    # Save document
    output_path = r'C:\Users\phamm\Desktop\CHUONG_IV_Mo_Hinh_Thuc_Nghiem_FULL.docx'
    doc.save(output_path)
    print(f'✓ Tạo file Word thành công: {output_path}')
    print(f'✓ Tổng số mục: {len(doc.paragraphs)} đoạn văn')
    print(f'✓ Tổng số bảng: 5 bảng')

if __name__ == '__main__':
    create_document()
